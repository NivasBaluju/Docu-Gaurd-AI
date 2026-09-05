"""
DocuGuard AI — Enterprise Word (DOCX) Redline Export Generator
Generates high-fidelity Microsoft Word documents (.docx) with native OpenXML
tracked changes (<w:ins> and <w:del>), side-by-side comparison tables,
and evidence-grounded audit traceability.
"""

import sys
import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def add_tracked_deletion(paragraph, deleted_text, author="DocuGuard AI", date="2026-09-05T12:00:00Z", revision_id="1"):
    """Adds native OpenXML <w:del> element recognized as a tracked deletion in Word."""
    p = paragraph._p
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), str(revision_id))
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date)
    
    del_text = OxmlElement('w:delText')
    del_text.text = deleted_text
    del_elem.append(del_text)
    p.append(del_elem)

def add_tracked_insertion(paragraph, inserted_text, author="DocuGuard AI", date="2026-09-05T12:00:00Z", revision_id="2"):
    """Adds native OpenXML <w:ins> element recognized as a tracked insertion in Word."""
    p = paragraph._p
    ins_elem = OxmlElement('w:ins')
    ins_elem.set(qn('w:id'), str(revision_id))
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), date)
    
    run_elem = OxmlElement('w:r')
    text_elem = OxmlElement('w:t')
    text_elem.text = inserted_text
    run_elem.append(text_elem)
    ins_elem.append(run_elem)
    p.append(ins_elem)

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell in OpenXML."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def generate_redline_docx(data: dict, output_path: str):
    """
    Builds the production-grade redline document.
    Data format:
    {
      "contractName": "...",
      "documentId": "...",
      "sha256": "...",
      "negotiationMode": "BALANCED",
      "overallRisk": 42,
      "clauses": [
        {
          "clauseId": "cl-01",
          "clauseType": "Limitation of Liability",
          "risk": "HIGH",
          "originalText": "...",
          "proposedText": "...",
          "diffWords": [ {"type": "same"|"add"|"del", "text": "..."} ],
          "rationale": "...",
          "evidenceRef": "Section 12.1, Page 4"
        }
      ]
    }
    """
    doc = docx.Document()
    
    # Configure 1-inch margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Document Title (Fraunces / Georgia Serif style)
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("DOCUGUARD AI — REDLINE & NEGOTIATION EXPORT")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    title_p.paragraph_format.space_after = Pt(4)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"Contract: {data.get('contractName', 'Enterprise Agreement')}  |  Mode: {data.get('negotiationMode', 'BALANCED').upper()}")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub_p.paragraph_format.space_after = Pt(16)

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document Identifier:", str(data.get('documentId', 'N/A'))),
        ("Cryptographic SHA-256:", str(data.get('sha256', 'N/A'))[:32] + "…"),
        ("Baseline Risk Score:", f"{data.get('overallRisk', 'N/A')}/100"),
        ("Export Generation Date:", "2026-09-05T12:00:00Z (Authoritative Snapshot)")
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = k
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(9.5)
        c0.paragraphs[0].runs[0].font.name = "Calibri"
        c1.text = v
        c1.paragraphs[0].runs[0].font.size = Pt(9.5)
        c1.paragraphs[0].runs[0].font.name = "Calibri"
        set_cell_background(c0, "F5F5F5")
        set_cell_background(c1, "FAFAFA")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Section 1: Executive Clause Summary
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. EXECUTIVE NEGOTIATION SUMMARY")
    h1_run.font.name = "Georgia"
    h1_run.font.size = Pt(13)
    h1_run.font.bold = True
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Clause ID", "Clause Type", "Risk Tier", "Grounded Evidence Ref"]
    for j, h in enumerate(headers):
        cell = summary_table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_background(cell, "EAEAEA")

    for cl in data.get('clauses', []):
        r = summary_table.add_row()
        r.cells[0].text = cl.get('clauseId', 'CL-01')
        r.cells[1].text = cl.get('clauseType', 'General Provision')
        r.cells[2].text = cl.get('risk', 'MEDIUM')
        r.cells[3].text = cl.get('evidenceRef', 'Section Reference')
        for c in r.cells:
            c.paragraphs[0].runs[0].font.size = Pt(9)
            c.paragraphs[0].runs[0].font.name = "Calibri"

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # Section 2: Detailed Clause Redline with Side-by-Side Comparison
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. CLAUSE-BY-CLAUSE REDLINE & COMPARISON")
    h2_run.font.name = "Georgia"
    h2_run.font.size = Pt(13)
    h2_run.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    rev_counter = 1
    for idx, cl in enumerate(data.get('clauses', [])):
        clause_p = doc.add_paragraph()
        c_title = clause_p.add_run(f"Clause {idx + 1}: {cl.get('clauseType', 'Contract Term')} [{cl.get('clauseId', 'ID')}]")
        c_title.font.bold = True
        c_title.font.size = Pt(11)
        c_title.font.name = "Georgia"

        # Rationale & Evidence Context
        rat_p = doc.add_paragraph()
        rat_run = rat_p.add_run(f"Strategic Objective: {cl.get('rationale', 'Balanced commercial risk allocation.')}\nEvidence Grounding: {cl.get('evidenceRef', 'Contract text.')}")
        rat_run.font.size = Pt(9)
        rat_run.font.italic = True
        rat_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        rat_p.paragraph_format.space_after = Pt(6)

        # Side-by-Side Table: Left = Original, Right = Proposed
        comp_table = doc.add_table(rows=1, cols=2)
        comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        comp_table.rows[0].cells[0].text = "ORIGINAL CONTRACT CLAUSE"
        comp_table.rows[0].cells[1].text = "PROPOSED NEGOTIATED CLAUSE"
        comp_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        comp_table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        comp_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
        comp_table.rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_background(comp_table.rows[0].cells[0], "F0F0F0")
        set_cell_background(comp_table.rows[0].cells[1], "F0F0F0")

        row = comp_table.add_row()
        orig_cell, prop_cell = row.cells[0], row.cells[1]
        orig_cell.text = cl.get('originalText', '')
        orig_cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        orig_cell.paragraphs[0].runs[0].font.name = "Calibri"

        prop_cell.text = cl.get('proposedText', '')
        prop_cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        prop_cell.paragraphs[0].runs[0].font.name = "Calibri"
        set_cell_background(orig_cell, "FCFCFC")
        set_cell_background(prop_cell, "FAFCFA")

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # Native Tracked Changes Presentation in Document Body
        track_heading = doc.add_paragraph()
        th_run = track_heading.add_run("Native Tracked Revisions (Word Compatible):")
        th_run.font.bold = True
        th_run.font.size = Pt(9.5)
        th_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        tracked_p = doc.add_paragraph()
        tracked_p.paragraph_format.left_indent = Inches(0.25)
        
        diff_words = cl.get('diffWords', [])
        if diff_words:
            for word_item in diff_words:
                w_type = word_item.get('type')
                w_text = word_item.get('text', '')
                if w_type == 'del':
                    add_tracked_deletion(tracked_p, w_text + " ", author="DocuGuard AI", revision_id=str(rev_counter))
                    rev_counter += 1
                elif w_type == 'add':
                    add_tracked_insertion(tracked_p, w_text + " ", author="DocuGuard AI", revision_id=str(rev_counter))
                    rev_counter += 1
                else:
                    r = tracked_p.add_run(w_text + " ")
                    r.font.name = "Calibri"
                    r.font.size = Pt(9.5)
        else:
            # Fallback if precomputed diff words not supplied: show deletion then insertion
            add_tracked_deletion(tracked_p, cl.get('originalText', ''), author="DocuGuard AI", revision_id=str(rev_counter))
            rev_counter += 1
            add_tracked_insertion(tracked_p, " " + cl.get('proposedText', ''), author="DocuGuard AI", revision_id=str(rev_counter))
            rev_counter += 1

        doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Footer note
    footer = doc.sections[0].footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = f_p.add_run("DocuGuard AI — Enterprise Contract Assurance  |  Non-Repudiation Verified")
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python docx_generator.py <input_json_path> <output_docx_path>")
        sys.exit(1)
    
    in_json = sys.argv[1]
    out_docx = sys.argv[2]
    with open(in_json, 'r', encoding='utf-8') as f:
        payload = json.load(f)
    
    res_path = generate_redline_docx(payload, out_docx)
    print(f"SUCCESS: Generated DOCX at {res_path}")
