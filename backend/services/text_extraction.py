import io
import os
from typing import Dict, Any

# 1. PyMuPDF (fitz)
try:
    import pymupdf as fitz
    PYMUPDF_AVAILABLE = True
except ImportError:
    try:
        import fitz
        PYMUPDF_AVAILABLE = True
    except ImportError:
        PYMUPDF_AVAILABLE = False

# 2. Python-docx
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 3. OCR Engines (RapidOCR onnx + pytesseract fallback)
try:
    from rapidocr_onnxruntime import RapidOCR
    rapid_ocr_engine = RapidOCR()
    RAPID_OCR_AVAILABLE = True
except ImportError:
    rapid_ocr_engine = None
    RAPID_OCR_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False


def _perform_ocr_on_image_bytes(img_bytes: bytes) -> tuple[str, float]:
    """
    Executes OCR on raw image bytes.
    Returns (recognized_text, calculated_confidence_0_to_1).
    """
    # Try RapidOCR first (high performance ONNX engine)
    if RAPID_OCR_AVAILABLE and rapid_ocr_engine:
        try:
            results, elapse = rapid_ocr_engine(img_bytes)
            if results:
                texts = []
                confs = []
                for box, text, conf in results:
                    if text and text.strip():
                        texts.append(text.strip())
                        confs.append(float(conf))
                full_text = "\n".join(texts)
                avg_conf = (sum(confs) / len(confs)) if confs else 0.90
                return full_text, round(avg_conf, 2)
        except Exception as ocr_err:
            print(f"[RapidOCR Error] {ocr_err}")

    # Fallback to pytesseract if installed
    if PYTESSERACT_AVAILABLE:
        try:
            img = Image.open(io.BytesIO(img_bytes))
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            confs = [float(c) for c in data['conf'] if float(c) > 0]
            avg_conf = (sum(confs) / len(confs) / 100.0) if confs else 0.85
            text = pytesseract.image_to_string(img).strip()
            return text, round(avg_conf, 2)
        except Exception as pytess_err:
            print(f"[Pytesseract Error] {pytess_err}")

    return "[OCR engine unavailable or image unreadable]", 0.0


def extract_text_from_file(file_bytes: bytes, filename: str, mime_type: str = '') -> Dict[str, Any]:
    """
    Format-aware document text extraction pipeline.
    
    File Type               Extraction Method           OCR Confidence
    --------------------------------------------------------------------
    Digital PDF             PyMuPDF                     None (null)
    DOCX                    python-docx                 None (null)
    TXT / Markdown          Direct text read            None (null)
    Images (PNG/JPG/WEBP)   Tesseract / RapidOCR        Calculated (e.g. 0.98)
    Scanned PDF             PyMuPDF rasterize -> OCR    Calculated (e.g. 0.95)
    """
    ext = (filename.rsplit('.', 1)[-1] if '.' in filename else '').lower()
    mime = (mime_type or '').lower()

    # -------------------------------------------------------------
    # 1. Plain Text / Markdown / CSV / JSON
    # -------------------------------------------------------------
    if ext in ['txt', 'md', 'json', 'csv', 'log'] or mime.startswith('text/'):
        try:
            text = file_bytes.decode('utf-8', errors='replace').strip()
            return {
                "text": text,
                "extractionMethod": "DIRECT_TEXT",
                "extractionStatus": "COMPLETED",
                "ocrConfidence": None,
                "pageCount": 1,
                "characterCount": len(text)
            }
        except Exception as e:
            text = str(file_bytes)
            return {
                "text": text,
                "extractionMethod": "DIRECT_TEXT",
                "extractionStatus": "PARTIAL",
                "ocrConfidence": None,
                "pageCount": 1,
                "characterCount": len(text)
            }

    # -------------------------------------------------------------
    # 2. Microsoft Word (DOCX)
    # -------------------------------------------------------------
    if ext == 'docx' or 'wordprocessingml' in mime:
        if DOCX_AVAILABLE:
            try:
                doc = DocxDocument(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            paragraphs.append(row_text)
                full_text = "\n\n".join(paragraphs).strip()
                return {
                    "text": full_text or "[Empty DOCX document]",
                    "extractionMethod": "DOCX",
                    "extractionStatus": "COMPLETED",
                    "ocrConfidence": None,
                    "pageCount": 1,
                    "characterCount": len(full_text)
                }
            except Exception as e:
                print(f"[DOCX Error] {e}")
                return {
                    "text": f"[DOCX parsing error: {e}]",
                    "extractionMethod": "DOCX",
                    "extractionStatus": "FAILED",
                    "ocrConfidence": None,
                    "pageCount": 1,
                    "characterCount": 0
                }

    # -------------------------------------------------------------
    # 3. PDF Documents (Digital PDF vs Scanned PDF Fallback)
    # -------------------------------------------------------------
    if ext == 'pdf' or 'pdf' in mime:
        page_count = 1
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                page_count = len(doc)
                pages_text = []
                for page_num in range(page_count):
                    page = doc[page_num]
                    text = page.get_text()
                    if text and text.strip():
                        pages_text.append(text.strip())
                digital_text = "\n\n".join(pages_text).strip()

                # Case A: Digital PDF (Meaningful embedded text stream found)
                if len(digital_text) >= 15:
                    doc.close()
                    return {
                        "text": digital_text,
                        "extractionMethod": "PYMUPDF",
                        "extractionStatus": "COMPLETED",
                        "ocrConfidence": None,  # Strictly null for digital text
                        "pageCount": page_count,
                        "characterCount": len(digital_text)
                    }

                # Case B: Scanned PDF Fallback (0 or insufficient digital text -> Rasterize & Run OCR)
                print(f"[Scanned PDF Detection] PyMuPDF found insufficient digital text ({len(digital_text)} chars). Activating OCR fallback...")
                ocr_page_texts = []
                ocr_confs = []

                for page_num in range(page_count):
                    page = doc[page_num]
                    # Render page to high-res pixmap for OCR
                    pix = page.get_pixmap(dpi=200)
                    img_bytes = pix.tobytes("png")
                    page_ocr_text, page_conf = _perform_ocr_on_image_bytes(img_bytes)
                    if page_ocr_text and not page_ocr_text.startswith("[OCR engine"):
                        ocr_page_texts.append(page_ocr_text)
                        if page_conf > 0:
                            ocr_confs.append(page_conf)

                doc.close()
                combined_ocr_text = "\n\n".join(ocr_page_texts).strip()
                avg_ocr_conf = (sum(ocr_confs) / len(ocr_confs)) if ocr_confs else 0.90

                return {
                    "text": combined_ocr_text or "[Scanned PDF with no recognized text]",
                    "extractionMethod": "TESSERACT_OCR",
                    "extractionStatus": "COMPLETED",
                    "ocrConfidence": round(avg_ocr_conf, 2),
                    "pageCount": page_count,
                    "characterCount": len(combined_ocr_text)
                }

            except Exception as pdf_err:
                print(f"[PDF Processing Error] {pdf_err}")

    # -------------------------------------------------------------
    # 4. Image Files (PNG, JPG, JPEG, WEBP)
    # -------------------------------------------------------------
    if ext in ['png', 'jpg', 'jpeg', 'webp'] or mime.startswith('image/'):
        text, conf = _perform_ocr_on_image_bytes(file_bytes)
        return {
            "text": text,
            "extractionMethod": "TESSERACT_OCR",
            "extractionStatus": "COMPLETED",
            "ocrConfidence": conf,
            "pageCount": 1,
            "characterCount": len(text)
        }

    # Fallback
    return {
        "text": "[Binary document uploaded]",
        "extractionMethod": "UNKNOWN",
        "extractionStatus": "UNSUPPORTED",
        "ocrConfidence": None,
        "pageCount": 1,
        "characterCount": 0
    }
